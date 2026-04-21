"""
Build Amazon KDP-compliant book artifacts for the 天机录 interactive novel
in 4 languages (zh, en, ja, ko).

Each language is published as a SEPARATE Amazon listing (separate ASIN), per
KDP rules — there is no automatic linking between language editions.

Inputs:
    output/天机录/if/chapters/ch0001.json … ch1200.json              (zh source)
    output/天机录/if/story_package.json                                (book bible)
    output/天机录/if/if_progress.json                                  (volume plan)
    output/天机录/translations/{en,ja,ko}/chapters/ch0001.json …      (translations)
    output/天机录/images/covers/cover_vol*.png                         (cover art)

Conversion strategy:
    Each interactive chapter is a sequence of text/dialogue/choice nodes.
    For every choice we pick option A — this is the canonical "扮猪吃虎"
    mainline path that matches the protagonist's persona — and inline its
    result_nodes. The choice prompt itself is dropped because it is meta
    text aimed at the reader-as-player.

Outputs (under output/天机录/amazon/{lang}/):
    docx/    — per-volume DOCX (10 vols × ~120 ch)
    epub/    — combined book EPUB
    metadata/ — book_description, keywords, categories, author_bio, book_metadata
    KDP_SUBMISSION_GUIDE.md
Plus a top-level output/天机录/amazon/MULTILINGUAL_RELEASE_PLAN.md.
"""

from __future__ import annotations

import json
import re
import shutil
import sys
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
from ebooklib import epub

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "output" / "天机录"
IF_DIR = SRC / "if"
ZH_CHAPTERS_DIR = IF_DIR / "chapters"
TRANS_DIR = SRC / "translations"
COVERS_DIR = SRC / "images" / "covers"
OUT_ROOT = SRC / "amazon"

# Per-language configuration. Source dir is required; everything else is
# the localized metadata that will appear on Amazon.
LANG_CONFIG: dict[str, dict] = {
    "zh": {
        "name": "Chinese (Simplified)",
        "kdp_language": "Chinese (Traditional) (Beta)",
        "locale": "zh-CN",
        "chapter_dir": ZH_CHAPTERS_DIR,
        "book_title": "天机录",
        "subtitle": "谋略修仙·扮猪吃虎·命运棋局",
        "author": "命书工作室",
        "chapter_label": lambda n, t: f"第{n}章 {t}",
        "volume_label": lambda i, t: f"卷{['一','二','三','四','五','六','七','八','九','十'][i]}：{t}",
        "volume_titles": ["废材觉醒","风云际会","秘境争锋","天机之谜","必死劫数","天眼降临","天机真相","命运终章","天界弈局","天命归一"],
        "synopsis": (
            "废柴弟子陈机意外获得残缺天机录，获得预见未来三时辰的能力，却每次使用都要消耗珍贵的天命值。"
            "他以废物之姿隐于众人之下，实则每一步都精准算计，在三大宗门倾轧、魔道入侵、上古禁地开启的乱世中，"
            "用片刻预见编织无人能破的棋局。但天机录越完整，命运的必死反噬就越强——他看得见所有人的结局，却始终看不透自己的命运。"
        ),
        "marketing_lines": [
            "★ 谋略修仙 · 扮猪吃虎 · 反转打脸 · 多线伏笔 ★",
            "★ 全书十卷·1200 章完整收录，跨越宗门暗战、秘境争锋、魔道倾轧、上古真相、天命弈局五重主题。",
            "★ 苏青瑶、夜清、韩烈、凌渊——朋友与仇敌交织的命运棋盘正待落子。",
        ],
        "keywords": ["修仙小说","扮猪吃虎","谋略升级","宗门争霸","穿越重生","玄幻奇幻","热血反转"],
        "categories": [
            "Books > Literature & Fiction > Genre Fiction > Coming of Age",
            "Books > Science Fiction & Fantasy > Fantasy > Epic",
            "Kindle eBooks > Foreign Languages > Chinese",
        ],
        "author_bio": "命书工作室专注东方玄幻与谋略修仙题材，致力于为读者打造结构严密、伏笔精巧的长篇网文。",
        "front_matter": {"title_page_label": "原创长篇玄幻"},
    },
    "en": {
        "name": "English",
        "kdp_language": "English",
        "locale": "en-US",
        "chapter_dir": TRANS_DIR / "en" / "chapters",
        "book_title": "The Heaven's Annal",
        "subtitle": "A Cultivation Saga of Foresight and Fate",
        "author": "Mingshu Studio",
        "chapter_label": lambda n, t: f"Chapter {n}: {t}",
        "volume_label": lambda i, t: f"Volume {['I','II','III','IV','V','VI','VII','VIII','IX','X'][i]}: {t}",
        "volume_titles": [
            "Awakening of the Discarded",
            "Storm of Gathering Winds",
            "Clash in the Secret Realm",
            "Riddle of Heaven's Will",
            "The Inescapable Doom",
            "Descent of the Heavenly Eye",
            "Truth Behind the Annal",
            "Final Chapter of Fate",
            "The Celestial Game",
            "Reunion of Destiny",
        ],
        "synopsis": (
            "Cast aside as a worthless disciple after his father's mysterious disappearance, "
            "Chen Ji unexpectedly inherits a fragment of the Heaven's Annal — an ancient artifact "
            "that grants him three hours of foresight, paid for in his own life force. Hiding behind "
            "the mask of a 'wasted talent', he weaves intricate plots inside warring sects, against "
            "demonic infiltration, and through the gates of forbidden ancient ruins. Yet the more "
            "complete the Annal becomes, the heavier its inevitable backlash grows — he can foresee "
            "every fate but his own."
        ),
        "marketing_lines": [
            "★ Xianxia · Cultivation · Strategic Genius MC · Slow-Burn Revenge ★",
            "★ Full 10-volume, 1,200-chapter epic — sect intrigue, demon invasions, cosmic chess.",
            "★ Allies and rivals — Su Qingyao, Ye Qing, Han Lie, Ling Yuan — converge on a single board.",
        ],
        "keywords": [
            "xianxia","cultivation novel","weak to strong","strategic protagonist",
            "chinese fantasy","epic fantasy series","face slapping",
        ],
        "categories": [
            "Books > Science Fiction & Fantasy > Fantasy > Epic",
            "Books > Literature & Fiction > Action & Adventure > Heroic Fantasy",
            "Kindle eBooks > Science Fiction & Fantasy > Fantasy > Asian Myths & Legends",
        ],
        "author_bio": (
            "Mingshu Studio is a collective specializing in long-form Eastern fantasy and "
            "strategic cultivation fiction, crafting tightly plotted epics with deep foreshadowing "
            "and intricate character webs."
        ),
        "front_matter": {"title_page_label": "An Original Epic Cultivation Saga"},
    },
    "ja": {
        "name": "Japanese",
        "kdp_language": "Japanese",
        "locale": "ja-JP",
        "chapter_dir": TRANS_DIR / "ja" / "chapters",
        "book_title": "天機録",
        "subtitle": "謀略修仙・偽装弱者・運命の棋局",
        "author": "命書スタジオ",
        "chapter_label": lambda n, t: f"第{n}章 {t}",
        "volume_label": lambda i, t: f"巻{['一','二','三','四','五','六','七','八','九','十'][i]}：{t}",
        "volume_titles": [
            "廃材覚醒","風雲集結","秘境争鋒","天機の謎","必死の劫数",
            "天眼降臨","天機の真相","運命終章","天界弈局","天命帰一",
        ],
        "synopsis": (
            "父の失踪後、宗門の「廃材」と蔑まれた弟子・陳機。彼は偶然、欠けた天機録を手にし、"
            "三時辰先の未来を視る力を得る——その代償は自身の天命値。廃物の仮面の下で全てを冷徹に"
            "計算しながら、三大宗門の暗闘、魔道の侵略、上古遺跡の解放という乱世を、断片的な"
            "予知を糸として誰にも破れぬ棋局に編んでいく。だが天機録が完全に近づくほど、"
            "「必死の反噬」も重くのしかかる——彼はあらゆる者の結末を視るが、自らの運命だけは見えない。"
        ),
        "marketing_lines": [
            "★ 仙侠 × 修仙 × 知略主人公 × 逆襲打顔 ★",
            "★ 全10巻・1200章の長編大作。宗門暗闘から天界の弈局まで五大主題を網羅。",
            "★ 蘇青瑶・夜清・韓烈・凌淵——盟友と仇敵が交錯する運命の棋盤、いざ最初の一手を。",
        ],
        "keywords": [
            "仙侠","修仙","異世界","東洋ファンタジー","謀略系主人公","逆襲","長編小説",
        ],
        "categories": [
            "本 > 文学・評論 > 中国・台湾文学",
            "Kindleストア > Kindle本 > 文学・評論 > ファンタジー",
            "Kindleストア > Kindle本 > SF・ホラー・ファンタジー > ファンタジー",
        ],
        "author_bio": (
            "命書スタジオは東洋ファンタジーと謀略修仙ジャンルに特化した制作集団。"
            "緻密な構成と伏線回収にこだわった長編作品を世に送り出している。"
        ),
        "front_matter": {"title_page_label": "オリジナル長編ファンタジー"},
    },
    "ko": {
        "name": "Korean",
        "kdp_language": "Korean (Beta)",
        "locale": "ko-KR",
        "chapter_dir": TRANS_DIR / "ko" / "chapters",
        "book_title": "천기록",
        "subtitle": "모략 수선 · 위장 약자 · 운명의 기국",
        "author": "명서 스튜디오",
        "chapter_label": lambda n, t: f"제{n}장 {t}",
        "volume_label": lambda i, t: f"제{['일','이','삼','사','오','육','칠','팔','구','십'][i]}권：{t}",
        "volume_titles": [
            "폐재 각성","풍운제회","비경 쟁봉","천기의 수수께끼","필사의 겁수",
            "천안 강림","천기의 진상","운명의 종장","천계 혁국","천명귀일",
        ],
        "synopsis": (
            "아버지의 실종 이후 종문에서 「폐물」로 멸시받던 천기. 그는 우연히 결손된 천기록을 손에 넣고, "
            "세 시진 앞의 미래를 예견하는 힘을 얻는다 — 단, 사용할 때마다 귀중한 천명값이 소모된다. "
            "폐물의 가면 아래에서 모든 수를 정밀하게 계산하며, 삼대 종문의 암투·마도의 침략·상고 금지의 개방이 "
            "교차하는 난세에서 짧은 예견을 실로 짜듯 누구도 풀 수 없는 기국을 짠다. "
            "그러나 천기록이 완전해질수록 「필사의 반噬」도 거세진다 — 그는 모두의 결말을 보지만, 자신의 운명만은 보지 못한다."
        ),
        "marketing_lines": [
            "★ 선협 × 수선 × 모략형 주인공 × 반전 통쾌 ★",
            "★ 전 10권·1200장 완간. 종문 암투부터 천계 기국까지 5대 주제 총망라.",
            "★ 소청요·야청·한열·능연 — 동지와 숙적이 교차하는 운명의 기반 위에서 첫 수가 놓인다.",
        ],
        "keywords": [
            "선협","수선","무협","동양 판타지","장편소설","모략","주인공이 강한",
        ],
        "categories": [
            "Kindle eBooks > Foreign Languages > Korean",
            "Kindle eBooks > Science Fiction & Fantasy > Fantasy > Epic",
        ],
        "author_bio": (
            "명서 스튜디오는 동양 판타지와 모략 수선 장르에 특화된 창작 집단으로, "
            "치밀한 구성과 정교한 복선으로 장편 대작을 선보입니다."
        ),
        "front_matter": {"title_page_label": "오리지널 장편 판타지"},
    },
}


def load_volume_plan() -> list[dict]:
    progress = json.loads((IF_DIR / "if_progress.json").read_text(encoding="utf-8"))
    out = []
    for idx in sorted(progress["volume_plans"].keys(), key=int):
        v = progress["volume_plans"][idx]
        out.append(
            {
                "index": int(idx),
                "theme": v.get("theme", ""),
                "start": v["chapter_range"]["start"],
                "end": v["chapter_range"]["end"],
            }
        )
    return out


def clean_text(text: str) -> str:
    """Normalize whitespace and stray quote marks left over from JSON content."""
    if not text:
        return ""
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = text.strip()
    if text.startswith('"') and text.count('"') % 2 == 1:
        text = text[1:].lstrip()
    return text


def walk_nodes(nodes: Iterable[dict], paragraphs: list[str]) -> None:
    """Flatten interactive nodes into a linear paragraph stream.

    Handles 6 structural variants seen in the dataset:
      1. {"text": {"content": "..."}}                 — wrapped text node
      2. {"dialogue": {"content": "..."}}             — wrapped dialogue node
      3. {"choice": {"choices": [...]}}               — wrapped choice node
      4. {"text": "label", "result_nodes": [...]}     — nested choice OPTION
      5. {"content": "...", "character_id": "..."}    — bare text/dialogue
      6. {"nodes": [...]} or {"choices": [...]}       — un-wrapped containers
    Choice points always pick option A (mainline canonical path).
    """
    for node in nodes:
        if not isinstance(node, dict):
            continue

        if "result_nodes" in node and "content" not in node:
            walk_nodes(node["result_nodes"], paragraphs)
            continue
        if "nodes" in node and "content" not in node:
            walk_nodes(node["nodes"], paragraphs)
            continue
        if "choices" in node and "choice" not in node:
            choices = node["choices"]
            if choices:
                walk_nodes(choices[0].get("result_nodes") or choices[0].get("nodes") or [], paragraphs)
            continue

        if "choice" in node and isinstance(node["choice"], dict):
            choices = node["choice"].get("choices", [])
            if choices:
                walk_nodes(choices[0].get("result_nodes") or choices[0].get("nodes") or [], paragraphs)
            continue

        if "text" in node and isinstance(node["text"], dict):
            content = clean_text(node["text"].get("content", ""))
            if content:
                paragraphs.append(content)
            continue
        if "dialogue" in node and isinstance(node["dialogue"], dict):
            content = clean_text(node["dialogue"].get("content", ""))
            if content:
                paragraphs.append(content)
            continue

        if "content" in node:
            content = clean_text(node.get("content", ""))
            if content:
                paragraphs.append(content)


def chapter_paragraphs(chapter_path: Path) -> tuple[str, list[str]]:
    data = json.loads(chapter_path.read_text(encoding="utf-8"))
    title = clean_text(data.get("title", "")).strip().strip('"')
    paragraphs: list[str] = []
    walk_nodes(data.get("nodes", []), paragraphs)
    final: list[str] = []
    for block in paragraphs:
        for para in re.split(r"\n\s*\n", block):
            para = para.strip()
            if para:
                final.append(para)
    return title, final


def build_volume_docx(cfg: dict, vol: dict, out_path: Path) -> int:
    doc = Document()
    style = doc.styles["Normal"]
    if cfg["locale"].startswith("zh"):
        style.font.name = "SimSun"
    elif cfg["locale"].startswith("ja"):
        style.font.name = "Yu Mincho"
    elif cfg["locale"].startswith("ko"):
        style.font.name = "Malgun Gothic"
    style.font.size = Pt(11)

    # Title page
    title_p = doc.add_paragraph()
    title_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_p.add_run(cfg["book_title"])
    title_run.bold = True
    title_run.font.size = Pt(28)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    vol_title = cfg["volume_titles"][vol["index"]]
    sub_run = sub_p.add_run(cfg["volume_label"](vol["index"], vol_title))
    sub_run.font.size = Pt(20)

    author_p = doc.add_paragraph()
    author_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author_p.add_run(cfg["author"]).font.size = Pt(14)

    if cfg["front_matter"].get("title_page_label"):
        label_p = doc.add_paragraph()
        label_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        label_p.add_run(cfg["front_matter"]["title_page_label"]).italic = True

    doc.add_page_break()

    chapter_count = 0
    for ch_num in range(vol["start"], vol["end"] + 1):
        ch_path = cfg["chapter_dir"] / f"ch{ch_num:04d}.json"
        if not ch_path.exists():
            continue
        title, paragraphs = chapter_paragraphs(ch_path)
        if not paragraphs:
            continue
        h = doc.add_heading(level=1)
        h_run = h.add_run(cfg["chapter_label"](ch_num, title))
        h_run.font.color.rgb = RGBColor(0, 0, 0)
        for p in paragraphs:
            doc.add_paragraph(p)
        doc.add_page_break()
        chapter_count += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return chapter_count


def build_epub(cfg: dict, lang_key: str, volumes: list[dict], out_path: Path) -> None:
    book = epub.EpubBook()
    book.set_identifier(f"tianjilu-{lang_key}-2026")
    book.set_title(cfg["book_title"])
    book.set_language(cfg["locale"])
    book.add_author(cfg["author"])
    book.add_metadata("DC", "description", cfg["synopsis"])
    for tag in cfg["keywords"]:
        book.add_metadata("DC", "subject", tag)

    cover_path = COVERS_DIR / "cover_vol1_awakening.png"
    if cover_path.exists():
        book.set_cover("cover.png", cover_path.read_bytes())

    spine: list = ["nav"]
    toc: list = []

    intro = epub.EpubHtml(title="Introduction", file_name="intro.xhtml", lang=cfg["locale"])
    intro.set_content(
        f"""<html><head><meta charset='utf-8'/><title>{cfg['book_title']}</title></head><body>
        <h1 style='text-align:center'>{cfg['book_title']}</h1>
        <h3 style='text-align:center'>{cfg['subtitle']}</h3>
        <h4 style='text-align:center'>{cfg['author']}</h4>
        <p>{cfg['synopsis']}</p>
        </body></html>"""
    )
    book.add_item(intro)
    spine.append(intro)

    for vol in volumes:
        vol_chapters: list = []
        for ch_num in range(vol["start"], vol["end"] + 1):
            ch_path = cfg["chapter_dir"] / f"ch{ch_num:04d}.json"
            if not ch_path.exists():
                continue
            ch_title, paragraphs = chapter_paragraphs(ch_path)
            if not paragraphs:
                continue
            html_body = "".join(f"<p>{p}</p>" for p in paragraphs)
            heading = cfg["chapter_label"](ch_num, ch_title)
            chapter = epub.EpubHtml(
                title=heading,
                file_name=f"ch{ch_num:04d}.xhtml",
                lang=cfg["locale"],
            )
            chapter.set_content(
                f"""<html><head><meta charset='utf-8'/><title>{heading}</title></head><body>
                <h2>{heading}</h2>
                {html_body}
                </body></html>"""
            )
            book.add_item(chapter)
            spine.append(chapter)
            vol_chapters.append(chapter)
        if vol_chapters:
            vol_label = cfg["volume_label"](vol["index"], cfg["volume_titles"][vol["index"]])
            toc.append((epub.Section(vol_label), vol_chapters))

    book.toc = tuple(toc)
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = spine

    out_path.parent.mkdir(parents=True, exist_ok=True)
    epub.write_epub(str(out_path), book, {})


def write_metadata_files(cfg: dict, volumes: list[dict], out_dir: Path) -> None:
    out_dir.mkdir(parents=True, exist_ok=True)

    description = cfg["synopsis"] + "\n\n" + "\n".join(cfg["marketing_lines"])
    (out_dir / "book_description.txt").write_text(description, encoding="utf-8")
    (out_dir / "keywords.txt").write_text("\n".join(cfg["keywords"]), encoding="utf-8")
    (out_dir / "categories.txt").write_text("\n".join(cfg["categories"]), encoding="utf-8")
    (out_dir / "author_bio.txt").write_text(f"{cfg['author']}\n\n{cfg['author_bio']}", encoding="utf-8")

    book_meta = {
        "title": cfg["book_title"],
        "subtitle": cfg["subtitle"],
        "author": cfg["author"],
        "language_label": cfg["name"],
        "kdp_language_setting": cfg["kdp_language"],
        "locale": cfg["locale"],
        "total_chapters": 1200,
        "volumes": [
            {
                "volume": v["index"] + 1,
                "title": cfg["volume_titles"][v["index"]],
                "chapter_start": v["start"],
                "chapter_end": v["end"],
                "theme": v["theme"],
            }
            for v in volumes
        ],
    }
    (out_dir / "book_metadata.json").write_text(
        json.dumps(book_meta, ensure_ascii=False, indent=2), encoding="utf-8"
    )


def write_submission_guide(cfg: dict, lang_key: str, volumes: list[dict], out_path: Path) -> None:
    table_rows = "\n".join(
        f"| {cfg['volume_label'](v['index'], cfg['volume_titles'][v['index']])} | "
        f"{v['start']}–{v['end']} | docx/{cfg['book_title']}-vol{v['index']+1}.docx |"
        for v in volumes
    )

    guide = f"""# 《{cfg['book_title']}》Amazon KDP Submission Guide ({cfg['name']})

## 1. File Inventory
- `epub/{cfg['book_title']}.epub` — full 1,200-chapter compendium (recommended upload).
- `docx/*.docx` — split by 10 volumes (~120 chapters each), suitable as separate Kindle eBooks.
- `metadata/` — copy text for KDP fields (description, keywords, categories, author bio).
- Cover candidates under `output/天机录/images/covers/`.

## 2. KDP Metadata Fields
| Field | Value |
|------|----|
| Book Title | {cfg['book_title']} |
| Subtitle | {cfg['subtitle']} |
| Author | {cfg['author']} |
| Language | **{cfg['kdp_language']}** |
| Locale | {cfg['locale']} |
| Description | see `metadata/book_description.txt` (≤4000 chars) |
| Keywords | see `metadata/keywords.txt` (≤7) |
| Categories | see `metadata/categories.txt` |
| Cover | 1600×2560 px JPG/TIFF (upscale existing PNG before submit) |
| ISBN | not required for eBook |

## 3. Volume Breakdown
| Volume | Chapters | DOCX |
|--------|----------|------|
{table_rows}

## 4. Submission Flow
1. Sign in at [https://kdp.amazon.com](https://kdp.amazon.com).
2. **Create a new eBook** — DO NOT add this as an "edition" of the Chinese book; KDP has no automatic linking between language editions, so each language gets its own ASIN.
3. Set *Language* to **{cfg['kdp_language']}** and enter the localized title/subtitle/author shown above.
4. Paste fields from `metadata/`.
5. Upload `epub/{cfg['book_title']}.epub` (preferred) or DOCX. Validate locally with [Kindle Previewer](https://kdp.amazon.com/help/topic/G202131170).
6. Upload cover.
7. Choose DRM and KDP Select preferences.
8. Set price (USD 2.99–9.99 maximizes royalty rate).
9. Publish — typically live within 24–72 hours.

## 5. Interactive → Linear Conversion Notes
- Source data is a branching interactive novel (2–3 choices per decision point).
- This release follows the **Option A mainline** (best matches the protagonist's
  "play weak, strike hard" persona).
- Stripped: choice prompts, stat changes, affinity deltas, choice UI text.
- Preserved: all narrative prose, dialogue, chapter titles, chapter hooks.
- Total: 1,200 chapters across 10 volumes.

## 6. Multi-language Notes
This title is also being released in 3 other languages — see
`output/天机录/amazon/MULTILINGUAL_RELEASE_PLAN.md` for the cross-language rollout plan.
"""
    out_path.write_text(guide, encoding="utf-8")


def build_for_language(cfg: dict, lang_key: str, volumes: list[dict]) -> dict:
    print(f"\n=== {lang_key.upper()} / {cfg['name']} ===")
    out_dir = OUT_ROOT / lang_key
    if out_dir.exists():
        shutil.rmtree(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    docx_dir = out_dir / "docx"
    total_ch = 0
    for vol in volumes:
        vol_label = vol["index"] + 1
        out_path = docx_dir / f"{cfg['book_title']}-vol{vol_label}-{cfg['volume_titles'][vol['index']]}.docx"
        n = build_volume_docx(cfg, vol, out_path)
        total_ch += n
        print(f"  vol{vol_label} {cfg['volume_titles'][vol['index']]:30s}  {n:4d} chapters")

    epub_path = out_dir / "epub" / f"{cfg['book_title']}.epub"
    build_epub(cfg, lang_key, volumes, epub_path)
    print(f"  EPUB → {epub_path.relative_to(SRC)}  ({epub_path.stat().st_size/1024/1024:.2f} MB)")

    write_metadata_files(cfg, volumes, out_dir / "metadata")
    write_submission_guide(cfg, lang_key, volumes, out_dir / "KDP_SUBMISSION_GUIDE.md")
    return {"chapters_in_docx": total_ch, "epub_size_mb": round(epub_path.stat().st_size / 1024 / 1024, 2)}


def write_multilingual_plan(stats: dict[str, dict]) -> None:
    rows = "\n".join(
        f"| {k.upper()} | {LANG_CONFIG[k]['name']} | {LANG_CONFIG[k]['kdp_language']} | "
        f"{LANG_CONFIG[k]['book_title']} | {stats[k]['chapters_in_docx']} | "
        f"{stats[k]['epub_size_mb']} MB |"
        for k in stats
    )

    plan = f"""# 《天机录》多语言上架总览 / Multilingual Release Plan

## 1. 单本 vs. 多本：KDP 规则
Amazon KDP **每种语言必须独立上架**，无法把多语言合并到同一本：
- 每个语言版本对应一个独立 ASIN、独立产品页、独立评论、独立销量。
- 自动算法**不会**关联不同语言的版本（不像「不同版本/精装/平装」会合并）。
- 标准做法：在标题或副标题里附上 "(English Edition)"、"(日本語版)"、"(한국어판)" 提高识别度。
- 如果想让读者在同一作者页看到所有语言版本，需要在 *Author Central* 把所有书绑定到同一作者档案。

## 2. 本次产出
| Lang | 名称 | KDP Language 字段 | 上架标题 | 章数 | EPUB 体积 |
|------|------|-------------------|----------|------|-----------|
{rows}

每个语言的产物位于 `amazon/{{lang}}/` 下：
- `epub/`、`docx/`、`metadata/`、`KDP_SUBMISSION_GUIDE.md`

## 3. 推荐上架顺序
1. **中文（简体）** —— 受众最大、原始语言、最先验证转换质量。
2. **English** —— Royal Road / WuxiaWorld 已培养出庞大的西方仙侠读者群。
3. **Japanese** —— なろう系长期对修仙/异世界题材接受度高。
4. **Korean** —— 韩国 KakaoPage 长篇修真用户基数大。

## 4. 跨语言协调要点
- **作者笔名统一**：建议在 *Author Central* 把 4 个语言的作者档案合并为一个，
  或保持「中文笔名 + 罗马音」一致（命书工作室 / Mingshu Studio / 命書スタジオ / 명서 스튜디오）。
- **封面统一**：复用同一组 10 卷封面，避免重做美术。仅替换书名字样的语言。
- **价格策略**：建议 USD 2.99–4.99，4 个语言保持等价，方便 ROI 分析。
- **关键词本地化**：每语言的 7 个关键词必须用目标语言的搜索习惯，参考各 `metadata/keywords.txt`。
- **A+ Content**：每个 ASIN 都要单独提交一份 A+ Content（图文营销页）。

## 5. 翻译质量提示
EN/JA/KO 翻译来自 MiniMax 自动翻译，存在以下已知问题：
- 偶发中文残留（特别是 KO 中混有汉字）；
- 角色名一致性已经过 `glossary.json` 校对，但部分场景描写仍可能机翻味较重；
- 上架前建议人工抽检前 30 章，定下润色标准后再扫尾。

## 6. 后续可拓展
- 繁体中文：可由简体繁体一键转换（OpenCC）后单独上架，覆盖港台市场。
- 越南语 / 泰语 / 印尼语：东南亚仙侠付费阅读市场上升期，可视 ROI 复制流程。
- 有声书 / Audible：完成首语言销量验证后，与 ACX 合作分发英文版有声书。
"""
    (OUT_ROOT / "MULTILINGUAL_RELEASE_PLAN.md").write_text(plan, encoding="utf-8")


def main() -> int:
    if OUT_ROOT.exists():
        shutil.rmtree(OUT_ROOT)
    OUT_ROOT.mkdir(parents=True, exist_ok=True)

    volumes = load_volume_plan()
    print(f"Volume plan loaded: {len(volumes)} volumes")

    stats: dict[str, dict] = {}
    for lang_key, cfg in LANG_CONFIG.items():
        if not cfg["chapter_dir"].exists():
            print(f"[skip] {lang_key}: chapter dir missing: {cfg['chapter_dir']}")
            continue
        stats[lang_key] = build_for_language(cfg, lang_key, volumes)

    # Top-level shared assets.
    cover_src = COVERS_DIR / "cover_vol1_awakening.png"
    if cover_src.exists():
        shutil.copy(cover_src, OUT_ROOT / "cover_vol1_preview.png")
    write_multilingual_plan(stats)

    print(f"\nDone. {len(stats)} languages built under {OUT_ROOT}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
