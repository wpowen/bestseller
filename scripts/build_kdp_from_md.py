#!/usr/bin/env python3
"""Build Amazon KDP-compliant artifacts from markdown chapter files.

Supports novels stored as chapter-NNN.md files (plain linear fiction,
not interactive). Auto-detects language from content, splits into
volumes of ~120 chapters, and produces DOCX + EPUB + metadata.

Usage:
    python scripts/build_kdp_from_md.py output/xianxia-upgrade
    python scripts/build_kdp_from_md.py output/female-no-cp-1776303225 \
        --title "风眼" --author "命书工作室"
    python scripts/build_kdp_from_md.py output/superhero-fiction-1776147970 \
        --lang en --chapters-per-volume 100
"""
from __future__ import annotations

import json
import re
import shutil
import sys
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
from ebooklib import epub


def detect_lang(text: str) -> str:
    cjk = len(re.findall(r"[\u4e00-\u9fff]", text))
    alpha = len(re.findall(r"[a-zA-Z]", text))
    if cjk > alpha:
        return "zh"
    return "en"


def parse_chapter_md(path: Path) -> tuple[str, list[str]]:
    raw = path.read_text(encoding="utf-8")
    lines = raw.split("\n")
    title = ""
    content_lines: list[str] = []
    for line in lines:
        if not title and line.startswith("#"):
            title = re.sub(r"^#+\s*", "", line).strip()
            continue
        content_lines.append(line)
    if not title:
        title = path.stem
    paragraphs: list[str] = []
    buf: list[str] = []
    for line in content_lines:
        if line.strip() == "":
            if buf:
                paragraphs.append("\n".join(buf).strip())
                buf = []
        else:
            buf.append(line)
    if buf:
        paragraphs.append("\n".join(buf).strip())
    return title, [p for p in paragraphs if p]


LANG_CONFIGS = {
    "zh": {
        "locale": "zh-CN",
        "kdp_language": "Chinese (Traditional) (Beta)",
        "font": "SimSun",
        "chapter_label": lambda n, t: f"第{n}章 {t}",
        "volume_label": lambda i, t: f"卷{['一','二','三','四','五','六','七','八','九','十','十一','十二','十三','十四','十五'][i]}：{t}",
        "default_author": "命书工作室",
    },
    "en": {
        "locale": "en-US",
        "kdp_language": "English",
        "font": "Times New Roman",
        "chapter_label": lambda n, t: f"Chapter {n}: {t}",
        "volume_label": lambda i, t: f"Volume {['I','II','III','IV','V','VI','VII','VIII','IX','X','XI','XII','XIII','XIV','XV'][i]}: {t}",
        "default_author": "Mingshu Studio",
    },
}


def _extract_title_text(ch_title: str, lang: str) -> str:
    if lang == "zh":
        for sep in ("：", ":"):
            if sep in ch_title:
                return ch_title.split(sep, 1)[1].strip()
        m = re.match(r"第\d+章\s*(.*)", ch_title)
        if m:
            return m.group(1).strip()
        return ch_title.strip()
    for sep in (":", "："):
        if sep in ch_title:
            return ch_title.split(sep, 1)[1].strip()
    m = re.match(r"Chapter\s+\d+\s*(.*)", ch_title, re.IGNORECASE)
    if m:
        return m.group(1).strip()
    return ch_title.strip()


def auto_volume_title(vol_idx: int, chapter_titles: list[str], lang: str) -> str:
    if not chapter_titles:
        return f"第{vol_idx+1}卷" if lang == "zh" else f"Volume {vol_idx + 1}"

    clean_titles = [_extract_title_text(t, lang) for t in chapter_titles if t]

    if lang == "zh":
        STOP = set("的了与及和在是不之一我他她它这那有大从到被让把给")
        bigram_freq: dict[str, int] = {}
        for t in clean_titles:
            filtered = "".join(ch for ch in t if ch not in STOP)
            for i in range(len(filtered) - 1):
                bg = filtered[i:i+2]
                bigram_freq[bg] = bigram_freq.get(bg, 0) + 1
        if bigram_freq:
            top = max(bigram_freq, key=bigram_freq.get)
            if bigram_freq[top] >= 2:
                return top
        for t in clean_titles:
            filtered = "".join(ch for ch in t if ch not in STOP)
            if len(filtered) >= 2:
                return filtered[:2]
        if clean_titles:
            return clean_titles[0][:2]
        return f"第{vol_idx+1}卷"

    if clean_titles:
        meta_kw = ("midpoint", "climax", "volume", "arc ", "part ", "act ",
                   "section ", "prologue", "epilogue", "interlude", "montage",
                   "recap", "end and beginning", "week ")
        def is_meta(t: str) -> bool:
            tl = t.lower()
            return any(kw in tl for kw in meta_kw) or any(tl.startswith(kw.strip()) for kw in meta_kw)
        good = [t for t in clean_titles if not is_meta(t)]
        pool = good if good else clean_titles
        candidates = sorted(pool, key=lambda t: (len(t.split()), len(t)))
        for c in reversed(candidates):
            words = c.split()
            if 2 <= len(words) <= 6 and len(c) <= 40:
                return c
        return candidates[-1][:30] if candidates else f"Volume {vol_idx + 1}"
    return f"Volume {vol_idx + 1}"


def build_volume_docx(
    cfg: dict,
    vol: dict,
    chapters: list[tuple[int, str, list[str]]],
    out_path: Path,
) -> int:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = cfg["font"]
    style.font.size = Pt(11)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_p.add_run(cfg["book_title"])
    title_run.bold = True
    title_run.font.size = Pt(28)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub_run = sub_p.add_run(cfg["volume_label"](vol["index"], vol["title"]))
    sub_run.font.size = Pt(20)

    author_p = doc.add_paragraph()
    author_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author_p.add_run(cfg["author"]).font.size = Pt(14)

    doc.add_page_break()

    count = 0
    for ch_num, ch_title, paragraphs in chapters:
        if not paragraphs:
            continue
        h = doc.add_heading(level=1)
        h_run = h.add_run(cfg["chapter_label"](ch_num, ch_title))
        h_run.font.color.rgb = RGBColor(0, 0, 0)
        for p in paragraphs:
            doc.add_paragraph(p)
        doc.add_page_break()
        count += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return count


def build_epub(
    cfg: dict,
    volumes: list[dict],
    all_chapters: dict[int, list[tuple[int, str, list[str]]]],
    out_path: Path,
) -> None:
    book = epub.EpubBook()
    book.set_identifier(f"{cfg['slug']}-{cfg['locale']}-2026")
    book.set_title(cfg["book_title"])
    book.set_language(cfg["locale"])
    book.add_author(cfg["author"])
    if cfg.get("synopsis"):
        book.add_metadata("DC", "description", cfg["synopsis"])

    spine: list = ["nav"]
    toc: list = []

    intro = epub.EpubHtml(
        title="Introduction", file_name="intro.xhtml", lang=cfg["locale"]
    )
    intro.set_content(
        f"""<html><head><meta charset='utf-8'/><title>{cfg['book_title']}</title></head><body>
        <h1 style='text-align:center'>{cfg['book_title']}</h1>
        <h3 style='text-align:center'>{cfg.get('subtitle','')}</h3>
        <h4 style='text-align:center'>{cfg['author']}</h4>
        <p>{cfg.get('synopsis','')}</p>
        </body></html>"""
    )
    book.add_item(intro)
    spine.append(intro)

    for vol in volumes:
        vol_chapters: list = []
        for ch_num, ch_title, paragraphs in all_chapters.get(vol["index"], []):
            if not paragraphs:
                continue
            html_body = "".join(f"<p>{p}</p>" for p in paragraphs)
            heading = cfg["chapter_label"](ch_num, ch_title)
            chapter = epub.EpubHtml(
                title=heading,
                file_name=f"ch{ch_num:04d}.xhtml",
                lang=cfg["locale"],
            )
            chapter.set_content(
                f"""<html><head><meta charset='utf-8'/><title>{heading}</title></head><body>
                <h2>{heading}</h2>
                {html_body}
                </body></html>"""
            )
            book.add_item(chapter)
            spine.append(chapter)
            vol_chapters.append(chapter)
        if vol_chapters:
            vol_label = cfg["volume_label"](vol["index"], vol["title"])
            toc.append((epub.Section(vol_label), vol_chapters))

    book.toc = tuple(toc)
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = spine

    out_path.parent.mkdir(parents=True, exist_ok=True)
    epub.write_epub(str(out_path), book, {})


def build_kdp(
    book_dir: Path,
    title: Optional[str] = None,
    author: Optional[str] = None,
    lang: Optional[str] = None,
    chapters_per_volume: int = 120,
    subtitle: str = "",
    synopsis: str = "",
    keywords: Optional[list[str]] = None,
    categories: Optional[list[str]] = None,
) -> dict:
    md_files = sorted(book_dir.glob("chapter-*.md"))
    if not md_files:
        print(f"[ERROR] No chapter-*.md files found in {book_dir}")
        return {}

    sample_text = md_files[0].read_text(encoding="utf-8")[:2000]
    detected_lang = lang or detect_lang(sample_text)
    lcfg = LANG_CONFIGS.get(detected_lang, LANG_CONFIGS["en"])

    slug = book_dir.name
    book_title = title or slug.replace("-", " ").title()
    book_author = author or lcfg["default_author"]

    print(f"=== {book_title} ({detected_lang.upper()}) ===")
    print(f"  Chapters: {len(md_files)}")
    print(f"  Author: {book_author}")

    all_parsed: list[tuple[int, str, list[str]]] = []
    for i, md in enumerate(md_files):
        ch_num = i + 1
        ch_title, paragraphs = parse_chapter_md(md)
        all_parsed.append((ch_num, ch_title, paragraphs))

    n_volumes = max(1, (len(all_parsed) + chapters_per_volume - 1) // chapters_per_volume)
    volumes: list[dict] = []
    vol_chapters: dict[int, list] = {}
    for vi in range(n_volumes):
        start = vi * chapters_per_volume
        end = min(start + chapters_per_volume, len(all_parsed))
        chunk = all_parsed[start:end]
        vol_ch_titles = [t for _, t, _ in chunk]
        vol_title = auto_volume_title(vi, vol_ch_titles, detected_lang)
        vol = {"index": vi, "title": vol_title, "start_ch": chunk[0][0] if chunk else 0, "end_ch": chunk[-1][0] if chunk else 0}
        volumes.append(vol)
        vol_chapters[vi] = chunk

    cfg = {
        "slug": slug,
        "book_title": book_title,
        "subtitle": subtitle,
        "author": book_author,
        "locale": lcfg["locale"],
        "kdp_language": lcfg["kdp_language"],
        "font": lcfg["font"],
        "chapter_label": lcfg["chapter_label"],
        "volume_label": lcfg["volume_label"],
        "synopsis": synopsis,
        "keywords": keywords or [],
        "categories": categories or [],
    }

    out_root = book_dir / "amazon" / detected_lang
    if out_root.exists():
        shutil.rmtree(out_root)
    out_root.mkdir(parents=True, exist_ok=True)

    docx_dir = out_root / "docx"
    total_ch = 0
    for vol in volumes:
        vol_label = vol["index"] + 1
        out_path = docx_dir / f"{book_title}-vol{vol_label}.docx"
        n = build_volume_docx(cfg, vol, vol_chapters[vol["index"]], out_path)
        total_ch += n
        print(f"  vol{vol_label} {vol['title']:30s} {n:4d} chapters")

    epub_path = out_root / "epub" / f"{book_title}.epub"
    build_epub(cfg, volumes, vol_chapters, epub_path)
    epub_mb = epub_path.stat().st_size / 1024 / 1024
    print(f"  EPUB → {epub_path.relative_to(book_dir)} ({epub_mb:.2f} MB)")

    meta_dir = out_root / "metadata"
    meta_dir.mkdir(parents=True, exist_ok=True)
    if synopsis:
        (meta_dir / "book_description.txt").write_text(synopsis, encoding="utf-8")
    if keywords:
        (meta_dir / "keywords.txt").write_text("\n".join(keywords), encoding="utf-8")
    if categories:
        (meta_dir / "categories.txt").write_text("\n".join(categories), encoding="utf-8")
    (meta_dir / "author_bio.txt").write_text(f"{book_author}\n", encoding="utf-8")

    book_meta = {
        "title": book_title,
        "subtitle": subtitle,
        "author": book_author,
        "language": detected_lang,
        "kdp_language_setting": lcfg["kdp_language"],
        "locale": lcfg["locale"],
        "total_chapters": len(all_parsed),
        "volumes": [
            {
                "volume": v["index"] + 1,
                "title": v["title"],
                "chapter_start": v["start_ch"],
                "chapter_end": v["end_ch"],
            }
            for v in volumes
        ],
    }
    (meta_dir / "book_metadata.json").write_text(
        json.dumps(book_meta, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    return {
        "slug": slug,
        "title": book_title,
        "lang": detected_lang,
        "chapters": total_ch,
        "volumes": n_volumes,
        "epub_mb": round(epub_mb, 2),
    }


BOOK_DEFINITIONS = {
    "xianxia-upgrade": {
        "title": "绝境修仙",
        "subtitle": "废柴崛起·步步为营·逆天改命",
        "keywords": ["修仙小说","废柴逆袭","宗门升级","玄幻奇幻","热血"],
        "categories": ["Kindle eBooks > Science Fiction & Fantasy > Fantasy > Epic"],
    },
    "exorcist-detective-1776865200": {
        "title": "双生血印",
        "subtitle": "风水悬疑·双生宿命·阴阳探秘",
        "keywords": ["风水悬疑","都市灵异","双生","推理","玄幻"],
        "categories": ["Kindle eBooks > Mystery & Thriller > Supernatural"],
    },
    "exorcist-detective-1778051012": {
        "title": "十五分钟凶宅",
        "subtitle": "风水铺·凶宅·十五分钟",
        "keywords": ["风水","凶宅","悬疑","都市灵异"],
        "categories": ["Kindle eBooks > Mystery & Thriller > Supernatural"],
    },
    "xianxia-upgrade-1776137730": {
        "title": "暗潮修仙录",
        "subtitle": "暗潮涌动·步步杀机·逆命登天",
        "keywords": ["修仙小说","暗潮","宗门权谋","玄幻奇幻","热血"],
        "categories": ["Kindle eBooks > Science Fiction & Fantasy > Fantasy > Epic"],
    },
    "female-no-cp-1776303225": {
        "title": "风眼",
        "subtitle": "都市异能·无CP·风眼露锋",
        "keywords": ["都市异能","无CP","女主","悬疑","热血"],
        "categories": ["Kindle eBooks > Science Fiction & Fantasy > Fantasy > Urban"],
    },
    "superhero-fiction-1776147970": {
        "title": "Cipher Threshold",
        "subtitle": "A Superhero Thriller of Power and Conspiracy",
        "lang": "en",
        "keywords": ["superhero fiction","urban fantasy","metahuman","conspiracy thriller","dark heroes"],
        "categories": ["Kindle eBooks > Science Fiction & Fantasy > Fantasy > Urban"],
    },
    "romantasy-1776330993": {
        "title": "Cinder Genesis",
        "subtitle": "A Romantasy of Ashes and Rebellion",
        "lang": "en",
        "keywords": ["romantasy","fantasy romance","fae courts","rebellion","enemies to lovers"],
        "categories": ["Kindle eBooks > Romance > Fantasy"],
    },
    "superhero-fiction-1776301343": {
        "title": "Ash Opening Move",
        "subtitle": "Superpowers, Survival, and the Price of Burning",
        "lang": "en",
        "keywords": ["superhero","urban fantasy","metahuman","action thriller","coming of age"],
        "categories": ["Kindle eBooks > Science Fiction & Fantasy > Fantasy > Urban"],
    },
}


def main() -> int:
    base = Path(__file__).resolve().parent.parent / "output"
    results: list[dict] = []

    for slug, defs in BOOK_DEFINITIONS.items():
        book_dir = base / slug
        if not book_dir.exists():
            print(f"[skip] {slug}: directory not found")
            continue
        md_files = list(book_dir.glob("chapter-*.md"))
        if not md_files:
            print(f"[skip] {slug}: no chapter markdown files")
            continue

        result = build_kdp(
            book_dir=book_dir,
            title=defs.get("title"),
            subtitle=defs.get("subtitle", ""),
            lang=defs.get("lang"),
            keywords=defs.get("keywords"),
            categories=defs.get("categories"),
        )
        if result:
            results.append(result)
        print()

    print("=" * 70)
    print(f"Built {len(results)} books:")
    for r in results:
        print(f"  {r['title']:30s} [{r['lang']}] {r['chapters']}ch / {r['volumes']}vol / {r['epub_mb']}MB epub")

    return 0


if __name__ == "__main__":
    sys.exit(main())
